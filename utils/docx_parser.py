import csv
from typing import List, Tuple, Dict
from docx import Document

def parse_handles_from_docx(path: str) -> List[Dict[str, str]]:
    doc = Document(path)
    items = []
    # Read tables
    for table in doc.tables:
        headers = [c.text.strip().lower() for c in table.rows[0].cells]
        for row in table.rows[1:]:
            record = {headers[i]: row.cells[i].text.strip() for i in range(len(headers))}
            platform = record.get("platform") or record.get("site") or ""
            handle = record.get("handle") or record.get("url") or record.get("profile") or ""
            if platform and handle:
                items.append({"platform": platform.lower(), "handle": handle})
    # Also scan paragraphs (one per line: platform,handle)
    for p in doc.paragraphs:
        line = p.text.strip()
        if "," in line:
            parts = [x.strip() for x in line.split(",")]
            if len(parts) >= 2:
                items.append({"platform": parts[0].lower(), "handle": parts[1]})
    return items

def parse_handles_from_csv(path: str) -> List[Dict[str, str]]:
    items = []
    with open(path, newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            items.append({"platform": row["platform"].lower().strip(), "handle": row["handle"].strip()})
    return items